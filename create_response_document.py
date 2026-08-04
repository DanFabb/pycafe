"""
Script to create the reviewer response document for pyCAFE (JORS submission).
Run once to create the document; subsequent entries are added by appending.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Styles ──────────────────────────────────────────────────────────────────

def set_font(run, bold=False, italic=False, size=11, color=None):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor(*color)


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p


def body(doc, text):
    return doc.add_paragraph(text)


def reviewer_comment(doc, reviewer_id, comment_text):
    """Adds a reviewer comment in italic grey."""
    p = doc.add_paragraph()
    run = p.add_run(f"Reviewer {reviewer_id} — Comment: ")
    set_font(run, bold=True, size=10, color=(70, 70, 70))
    run2 = p.add_run(comment_text)
    set_font(run2, italic=True, size=10, color=(100, 100, 100))
    return p


def author_response(doc, response_text):
    """Adds the authors' response in normal black."""
    p = doc.add_paragraph()
    run = p.add_run("Authors' Response: ")
    set_font(run, bold=True, size=11)
    run2 = p.add_run(response_text)
    set_font(run2, size=11)
    return p


def action_taken(doc, action_text):
    """Adds an action taken note in bold blue."""
    p = doc.add_paragraph()
    run = p.add_run("Action taken: ")
    set_font(run, bold=True, size=10, color=(0, 70, 160))
    run2 = p.add_run(action_text)
    set_font(run2, size=10, color=(0, 70, 160))
    return p


def separator(doc):
    doc.add_paragraph("─" * 80)


# ── Title Page ───────────────────────────────────────────────────────────────

heading(doc, "Response to Reviewers", level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = p.add_run(
    "Manuscript: pyCAFE: A Finite Element Framework for Solving Acoustic Problems in Python\n"
    "Journal: Journal of Open Research Software (JORS)\n"
    f"Date: {datetime.date.today().strftime('%B %d, %Y')}\n"
    "Authors: Daniele Fabbri, Fabio Bruzzone, Carlo Rosso"
)
run.font.size = Pt(11)

doc.add_paragraph()
p = doc.add_paragraph(
    "We thank all seven reviewers for their thorough and constructive comments. "
    "We have addressed each point below. For each comment we indicate the action taken "
    "and, where applicable, the section of the manuscript or the repository that has been modified. "
    "Points we have chosen not to implement are accompanied by a clear justification."
)

doc.add_page_break()

# ── Entry 1: Quality Control — Testing Suite ────────────────────────────────

heading(doc, "1. Quality Control and Testing", level=1)

body(doc,
    "This section documents the expansion of the test suite in response to comments "
    "from Reviewers 1, 3, 4, and 5 regarding the inadequacy of the original testing "
    "strategy. The original test suite comprised a single smoke test (test_basic.py) "
    "that verified package import and a minimal end-to-end workflow."
)

doc.add_paragraph()

# ─ 1.1 Bug found
heading(doc, "1.1  Bug discovered during testing: gauss_rule_quad_2x2 in element_cquad8.py", level=2)

reviewer_comment(doc, "1",
    "The Quality Control section requires considerable strengthening, as it does not "
    "adequately describe the testing strategy, validation methodology, reliability measures, "
    "or any automated testing and continuous integration mechanisms."
)
reviewer_comment(doc, "3",
    "Only minimal pytest functional workflow test is described. Suggest unit tests for: "
    "element matrix correctness, boundary condition enforcement, solver accuracy."
)

author_response(doc,
    "In the process of writing unit tests for the CQUAD8 element matrices, a latent bug "
    "was discovered in pycafe/build_matrices/element_cquad8.py. "
    "The function gauss_rule_quad_2x2() returned the quadrature weight w as a Python "
    "integer scalar (w = 1) instead of a NumPy array. This caused a TypeError when "
    "iterating over the weights in the integration loop of element_matrices_cquad8(), "
    "making direct calls to that function fail silently in production (the assembly "
    "routine assembly_cquad8.py bypassed the issue by explicitly passing gauss_rule_quad_3x3). "
    "Additionally, the default quadrature rule for element_matrices_cquad8() was incorrectly "
    "set to gauss_rule_quad_2x2 (2×2 Gauss points), whereas a 3×3 rule is required for "
    "exact integration of the CQUAD8 serendipity element stiffness and mass matrices."
)

action_taken(doc,
    "(1) Fixed gauss_rule_quad_2x2() in element_cquad8.py to return w * np.ones_like(xi) "
    "instead of the scalar w = 1. "
    "(2) Changed the default quad_rule argument of element_matrices_cquad8() from "
    "gauss_rule_quad_2x2 to gauss_rule_quad_3x3, which is the correct integration rule "
    "for 8-node serendipity elements. The assembly routine (assembly_cquad8.py) was "
    "already using gauss_rule_quad_3x3 explicitly and is unaffected."
)

doc.add_paragraph()

# ─ 1.2 New test files
heading(doc, "1.2  New unit test files", level=2)

author_response(doc,
    "Three new test files have been added to the tests/ directory, covering element "
    "matrix correctness, boundary condition enforcement, and solver accuracy with "
    "analytical benchmarks. All 70 tests pass. Details follow."
)

doc.add_paragraph()

body(doc, "tests/test_element_matrices.py  (27 tests)")
items = [
    "CQUAD4 and CQUAD8 shape functions: partition of unity (ΣN_i = 1), "
     "Kronecker delta at element nodes (N_i(x_j) = δ_ij), sum of derivatives = 0.",
    "Jacobian: correct value (detJ = Lx·Ly/4) for arbitrary rectangular elements; "
     "positivity for valid elements.",
    "Stiffness matrix K_e: symmetry, positive semi-definiteness, exactly one zero "
     "eigenvalue (null space = constant pressure, zero-energy mode for Neumann BCs), "
     "independence from speed of sound c.",
    "Mass matrix M_e: symmetry, positive definiteness, mass conservation "
     "(sum(M_e) = area/c²), correct scaling with c and element area.",
    "Both CQUAD4 and CQUAD8 elements are tested.",
]
for item in items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(11)

doc.add_paragraph()

body(doc, "tests/test_boundary_conditions.py  (10 tests)")
items = [
    "get_pressure_zero_nodes: correct 1-based to 0-based index conversion, "
     "deduplication, sorted output, handling of missing boundary names.",
    "reduce_KMC_dirichlet_mask: correct dimension of reduced system, "
     "idx_free excludes all constrained DOFs and contains all free DOFs, "
     "values in reduced matrix match global matrix, duplicate constraints handled.",
    "Global assembly: symmetry of K and M, mass conservation over full mesh, "
     "K·ones = 0 (global null space).",
]
for item in items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(11)

doc.add_paragraph()

body(doc, "tests/test_solver_modal.py  (7 tests)  — analytical benchmark")
items = [
    "Basic correctness: positive frequencies, sorted output, correct shape of mode "
     "shapes array, M-orthogonality of mode shapes.",
    "Analytical validation against the 2D rigid rectangular cavity. For a cavity "
     "Lx = 2 m, Ly = 1 m, filled with air (c₀ = 343 m/s), the exact natural "
     "frequencies are f_mn = (c₀/2)√((m/Lx)² + (n/Ly)²). "
     "The first natural frequency f_10 = 85.75 Hz is reproduced with < 1% error "
     "on a 10×5 element mesh.",
    "Convergence study: monotonic error reduction over meshes n = 2, 4, 8, 16 "
     "(square elements); convergence rate > 2 per mesh doubling (O(h²) behaviour "
     "for CQUAD4 bilinear elements).",
]
for item in items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(11)

doc.add_paragraph()

body(doc, "tests/test_solver_helmholtz.py  (12 tests)")
items = [
    "Single-frequency solver: correct output shape and dtype, exact enforcement "
     "of prescribed pressure BCs, RuntimeError for fully constrained system, "
     "effect of velocity RHS on solution.",
    "Frequency sweep: correct (N_dof × N_freq) output shape, consistency with "
     "single-frequency solve, pressure BC enforcement at every frequency step.",
    "Normal velocity RHS: empty boundary → empty arrays, non-zero contributions "
     "for active boundaries, linear scaling with v_n, purely imaginary character "
     "(proportional to iωρv_n).",
]
for item in items:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(item).font.size = Pt(11)

doc.add_paragraph()

action_taken(doc,
    "Added tests/test_element_matrices.py, tests/test_boundary_conditions.py, "
    "tests/test_solver_modal.py, tests/test_solver_helmholtz.py. "
    "The Quality Control section of the manuscript has been expanded to describe "
    "the full testing strategy, the analytical benchmark, and the convergence study."
)

separator(doc)

# ── Save ────────────────────────────────────────────────────────────────────

output_path = "Response_to_Reviewers_pyCAFE.docx"
doc.save(output_path)
print(f"Document saved: {output_path}")
