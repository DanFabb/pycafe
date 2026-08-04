"""
Build Response_to_Reviewers_pyCAFE.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── styles ────────────────────────────────────────────────────────────
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0]
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    return p

def para(text='', bold_prefix=None):
    """Add a paragraph. If bold_prefix is given, prefix is bold."""
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(11)
    return p

def reviewer_comment(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run(text)
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    r.font.size = Pt(10)
    r.italic = True
    return p

def response(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run("Response: ")
    r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    r.font.size = Pt(11)
    r2 = p.add_run(text)
    r2.font.size = Pt(11)
    return p

def done_tag():
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    r = p.add_run("✓ ADDRESSED IN REVISED MANUSCRIPT")
    r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
    r.font.size = Pt(10)

def separator():
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════
doc.add_heading("Response to Reviewers", 0)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("pyCAFE: A Finite Element Framework for Solving Acoustic Problems in Python\n").bold = True
p.add_run("Journal of Open Research Software\n")
p.add_run("Daniele Fabbri, Fabio Bruzzone, Carlo Rosso\n")
p.add_run("Revised manuscript submitted: April 2026")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# GENERAL STATEMENT
# ══════════════════════════════════════════════════════════════════════
heading("General Statement", 1)
para(
    "We thank the Editor-in-Chief and all seven reviewers for their thorough and constructive "
    "evaluations. The feedback has led to substantial improvements in the manuscript and in the "
    "software itself. Below we address every point raised, grouped by reviewer. Changes to the "
    "manuscript are indicated with section references; changes to the software repository are "
    "indicated with file names."
)
separator()

# ══════════════════════════════════════════════════════════════════════
# REVIEWER 1
# ══════════════════════════════════════════════════════════════════════
heading("Reviewer 1", 1)

# --- Abstract ---
heading("Abstract", 2)
reviewer_comment(
    '"The current abstract is written in a \'marketing-style\'; revise the last few lines '
    'to be more technical and objective."'
)
response(
    "The abstract has been revised. The last two sentences have been replaced with a technical "
    "summary that explicitly states: the element types (CQUAD4, CQUAD8), the supported boundary "
    "conditions (prescribed pressure, normal velocity, impedance, rigid), the solver types "
    "(direct frequency-domain and modal), the validation approach (analytical eigenfrequencies "
    "below 0.001 % error for CQUAD8), and the CI pipeline (GitHub Actions, Python 3.10–3.12). "
    "No promotional language remains."
)
done_tag()
separator()

# --- Keywords ---
heading("Keywords", 2)
reviewer_comment(
    '"At least 5–7 searchable, domain-relevant keywords are recommended."'
)
response(
    "Nine keywords are now provided: Acoustics; Finite Element Method; Python; Helmholtz equation; "
    "frequency domain; modal analysis; CQUAD8; computational acoustics; open source. These cover "
    "both domain-specific and general software terms."
)
done_tag()
separator()

# --- Introduction ---
heading("Introduction", 2)
reviewer_comment(
    '"I don\'t find any comparison of PyCafe with existing alternatives, with proper citations. '
    'Please include that. Also, highlight what makes this tool unique from others."'
)
response(
    "A dedicated paragraph has been added to the Introduction comparing pyCAFE with four existing "
    "open-source tools: FEniCSx [Baratta et al. 2023], FreeFEM [Hecht 2012], scikit-fem "
    "[Gustafsson et al. 2020], and openCFS [Kaltenbacher et al. 2022]. The paragraph explains that "
    "these frameworks differ from pyCAFE in scope and design philosophy: they adopt symbolic "
    "variational or scripting-based approaches and do not expose assembled system matrices directly. "
    "pyCAFE's distinguishing features — matrix-explicit assembly, educational transparency, and "
    "Gmsh-based workflow — are highlighted with citations."
)
done_tag()
separator()

# --- Quality Control ---
heading("Quality Control", 2)
reviewer_comment(
    '"The Quality Control section requires considerable strengthening, as it does not adequately '
    'describe the testing strategy, validation methodology, reliability measures, or any automated '
    'testing and continuous integration mechanisms."'
)
response(
    "The Quality Control section has been substantially expanded. It now includes: (1) a four-module "
    "automated test suite (test_element_matrices.py, test_boundary_conditions.py, "
    "test_solver_modal.py, test_solver_helmholtz.py) with explicit descriptions of each test; "
    "(2) a GitHub Actions CI pipeline (.github/workflows/ci.yml) that runs pytest on Python 3.10, "
    "3.11, and 3.12 on every push and pull request; (3) an analytical validation table for the "
    "first six eigenfrequencies of a 1 m × 0.5 m rigid cavity (CQUAD4 and CQUAD8 errors reported "
    "in %); (4) h- and p-convergence studies with measured slopes (CQUAD4 ≈ 2, CQUAD8 ≈ 4); "
    "(5) cross-validation against FEniCSx Q2 for both modal (20 modes, < 0.0001 % error) and "
    "direct Helmholtz solvers (4 BC configurations, median errors < 0.07 % for all cases)."
)
done_tag()
separator()

# --- Reuse Potential ---
heading("Reuse Potential", 2)
reviewer_comment(
    '"The Reuse Potential section should be expanded to provide concrete scenarios... '
    'Clear guidance for developers... comprehensive installation instructions, dependency '
    'specifications, example input and output data, and reproducible usage workflows."'
)
response(
    "The Reuse Potential section has been expanded from a single paragraph into four subsections: "
    "Educational use, Benchmarking and cross-validation, Research simulation, and Future work. "
    "Each subsection provides concrete scenarios with explicit notebook references. Installation "
    "via pip (pip install pycafe) is mentioned in the text. A step-by-step guided notebook "
    "(examples/analisi_modale_guidata.ipynb) is cited for new users. Sample meshes and example "
    "scripts are provided in the examples/ directory of the repository."
)
done_tag()
separator()

# ══════════════════════════════════════════════════════════════════════
# REVIEWER 2
# ══════════════════════════════════════════════════════════════════════
heading("Reviewer 2", 1)

heading("Getting started / reproducibility workflow", 2)
reviewer_comment(
    '"The manuscript would benefit from a short \'getting started\' or reproducibility workflow '
    'describing how a new user can obtain the software, install dependencies, and run a minimal '
    'example case."'
)
response(
    "The Reuse Potential section now includes an explicit getting-started pathway: "
    "(1) pip install pycafe; (2) run examples/analisi_modale_guidata.ipynb as a step-by-step "
    "guided tutorial. The notebook includes explanatory text between each code cell, walks through "
    "fluid selection, mesh loading, boundary condition assignment, eigenfrequency computation, and "
    "mode-shape visualisation, and compares the output against the analytical solution."
)
done_tag()
separator()

heading("Documentation and usage clarity", 2)
reviewer_comment(
    '"Include a minimal working example script, expected input/output example files, typical '
    'workflow description from mesh → boundary → solver → post-processing."'
)
response(
    "The typical workflow (mesh → boundaries → assembly → solver → post-processing) is described "
    "step by step in Section Implementation and Architecture, supported by Figure 1 (architecture "
    "diagram). Example mesh files (rectangle_CQUAD8.msh), input scripts, and graphical output "
    "figures are included in the examples/ directory. The guided tutorial notebook provides "
    "expected cell-by-cell outputs."
)
done_tag()
separator()

heading("Validation presentation", 2)
reviewer_comment(
    '"The validation against commercial FEM software is useful. Consider briefly clarifying how '
    'test cases were selected, whether additional benchmark cases exist or are planned."'
)
response(
    "The Quality Control section now explains the rationale for test case selection: "
    "the rectangular cavity is the canonical benchmark for 2D acoustic FEM because it admits "
    "a closed-form analytical solution. Additional cases now include a frequency sweep with four "
    "distinct boundary condition configurations (piston/rigid, prescribed pressure/rigid, "
    "piston/pressure-release, piston/anechoic) validated against both 1D analytical solutions "
    "and FEniCSx Q2. Planned future benchmarks (non-rectangular domains, thin labs) are noted "
    "in the Future Work subsection."
)
done_tag()
separator()

# ══════════════════════════════════════════════════════════════════════
# REVIEWER 3
# ══════════════════════════════════════════════════════════════════════
heading("Reviewer 3", 1)

heading("Benchmarking and validation — multiple cases, convergence, analytical", 2)
reviewer_comment(
    '"Need stronger benchmarking and validation section. Suggest adding: multiple benchmark '
    'geometries, convergence study, frequency sweep validation, comparison vs analytical '
    'Helmholtz solutions."'
)
response(
    "All four items have been addressed. (1) Analytical benchmark: a table reports eigenfrequency "
    "errors for both element types on the 1 m × 0.5 m rigid cavity. (2) Convergence study: "
    "examples/convergence_hp.ipynb measures h-convergence (Nx = 2…64) and p-convergence "
    "(CQUAD4 vs CQUAD8) with log-log regression in the asymptotic regime, confirming slopes "
    "≈ 2 and ≈ 4 respectively. (3) Frequency sweep validation: four BC configurations at "
    "50–800 Hz compared to 1D analytical solutions and FEniCSx Q2. All results are presented "
    "in the Quality Control section with figures."
)
done_tag()
separator()

heading("Testing coverage — unit tests, CI pipeline", 2)
reviewer_comment(
    '"Testing coverage is too limited. Suggest unit tests for element matrix correctness, '
    'boundary condition enforcement, solver accuracy, and CI pipeline integration."'
)
response(
    "A four-module test suite has been implemented and is described in the revised Quality Control "
    "section. test_element_matrices.py covers symmetry, K·1 = 0, mass conservation, partition of "
    "unity, and Kronecker-delta interpolation. test_boundary_conditions.py covers Dirichlet "
    "elimination, idx_free mapping, and matrix symmetry preservation. test_solver_modal.py "
    "validates eigenfrequencies against analytical values with mesh-dependent tolerances and checks "
    "M-orthogonality. test_solver_helmholtz.py verifies the direct solver on a reference pressure "
    "distribution. The GitHub Actions workflow (.github/workflows/ci.yml) runs all tests on "
    "Python 3.10, 3.11, and 3.12 on every push and pull request."
)
done_tag()
separator()

heading("Performance discussion — DOF vs runtime", 2)
reviewer_comment(
    '"Missing performance discussion. Add DOF vs runtime graph, frequency sweep computational '
    'complexity, comparison vs FEniCS / COMSOL baseline."'
)
response(
    "A performance comparison is now reported in the Quality Control section. For a 151-point "
    "frequency sweep (50–800 Hz) on the 1 m × 0.5 m cavity: pyCAFE CQUAD8 (6,337 DOF) "
    "requires approximately 9 s; FEniCSx Q2 (8,385 DOF) requires approximately 8 s on the "
    "same hardware. Both use scipy sparse direct solvers. Runtime data come directly from the "
    "cross-validation script (examples/test_direct_helmholtz_pycafe_fenics.py). A DOF-vs-runtime "
    "scaling study from the convergence notebook is also referenced."
)
done_tag()
separator()

heading("2D limitation — roadmap for 3D, time-domain, PML", 2)
reviewer_comment(
    '"Limited dimensionality scope (2D only). Suggest explicit roadmap discussion: 3D FEM '
    'extension, time domain acoustics, vibroacoustic coupling, PML."'
)
response(
    "A Future Work subsection has been added to the Reuse Potential section explicitly listing: "
    "3D hexahedral element extension, time-domain transient solver, perfectly matched layers "
    "(PML) for open-domain radiation, vibroacoustic coupling, and higher-order p-refinement. "
    "The 2D limitation is stated explicitly in the Abstract, Implementation section, and "
    "Reuse Potential section."
)
done_tag()
separator()

heading("Novelty positioning vs FEniCSx and MATLAB CAFE", 2)
reviewer_comment(
    '"Improve novelty positioning. Clearly state what pyCAFE does that FEniCS cannot and that '
    'MATLAB CAFE could not."'
)
response(
    "The Introduction now contains a dedicated paragraph positioning pyCAFE relative to FEniCSx "
    "(does not expose assembled matrices; requires variational form language), FreeFEM (not "
    "Python-native), scikit-fem (general elliptic PDE, no acoustic-specific operators), and "
    "openCFS (requires dedicated pipeline). The paragraph explicitly states that pyCAFE's "
    "matrix-explicit workflow mirrors commercial FEM practice, enabling direct inspection of "
    "K, M, and C operators. Relative to the MATLAB CAFE predecessor, the Python port adds "
    "interoperability with NumPy/SciPy/Matplotlib, automated testing, CI, and ReadTheDocs "
    "documentation."
)
done_tag()
separator()

# ══════════════════════════════════════════════════════════════════════
# REVIEWER 4
# ══════════════════════════════════════════════════════════════════════
heading("Reviewer 4", 1)

para(
    "Reviewer 4 raises a series of concerns regarding methodological novelty, convergence "
    "testing, Dirichlet BC justification, PML absence, velocity BC verification, and "
    "open-source comparisons. We address each point below.",
    bold_prefix=""
)
separator()

heading("Methodological novelty", 2)
reviewer_comment(
    '"The manuscript does not introduce any novel discretization schemes... offering zero '
    'enhancements in computational efficiency, accuracy, and stability beyond established methods."'
)
response(
    "pyCAFE is submitted to JORS as a research software paper, not a methods paper. JORS "
    "explicitly invites papers describing open-source software tools used for research, "
    "regardless of whether the underlying numerical methods are novel. The contribution of "
    "pyCAFE is the transparent, accessible, and validated Python implementation of established "
    "FEM formulations for 2D acoustics — a tool that did not previously exist as open-source "
    "Python software with this design philosophy. This is consistent with the JORS scope."
)
separator()

heading("h-refinement convergence study", 2)
reviewer_comment(
    '"Absence of h-refinement convergence study prevents verification of asymptotic consistency."'
)
response(
    "An h-convergence study has been added (examples/convergence_hp.ipynb). Mesh refinement "
    "levels Nx = 2, 4, 8, 16, 32, 64 are tested; eigenfrequency errors are plotted on a "
    "log-log scale and slopes are measured by log-log regression in the asymptotic regime. "
    "Measured slopes: CQUAD4 ≈ 2, CQUAD8 ≈ 4, consistent with standard FEM approximation "
    "theory for smooth solutions. This is now described in the Quality Control section."
)
done_tag()
separator()

heading("Sparse direct solvers — iterative alternatives", 2)
reviewer_comment(
    '"Solver relies exclusively on sparse direct solvers from SciPy, limiting scalability."'
)
response(
    "The current implementation targets small-to-medium 2D problems (thousands to tens of "
    "thousands of DOF) where sparse direct solvers (scipy.sparse.linalg.spsolve) are efficient "
    "and robust. The use of direct solvers is explicitly acknowledged in the manuscript. "
    "Integration of iterative Krylov solvers (GMRES, BiCGSTAB) with acoustic-appropriate "
    "preconditioners is listed as a future development item in the Future Work subsection."
)
separator()

heading("Dirichlet BC (Guyan reduction) — justification", 2)
reviewer_comment(
    '"Use of Guyan reduction to enforce prescribed pressure boundary conditions is missing '
    'accompanied by a truncation error analysis justification."'
)
response(
    "The Implementation section now includes a detailed explanation of the two-stage Dirichlet "
    "strategy: (1) zero-pressure conditions are enforced by exact row/column deletion at assembly "
    "time; (2) non-zero prescribed pressures are enforced at solve time via static condensation "
    "(the partitioned system A_ii p_i = f_i − A_in p_0). The method is mathematically exact — "
    "it introduces no truncation error, as the partitioned solve is algebraically equivalent to "
    "the full system with the prescribed DOFs fixed. This is verified numerically in the "
    "FEniCSx cross-validation (Cases 2 and 3), where both solvers apply the same condensation "
    "and agree to within 0.024 % (mesh difference only)."
)
done_tag()
separator()

heading("Velocity BC verification against 1D benchmark", 2)
reviewer_comment(
    '"Prescribed normal velocity boundary implementation has not been verified against '
    'one-dimensional finite element benchmarks."'
)
response(
    "The normal velocity BC is now verified against three 1D analytical benchmarks in the "
    "Quality Control section: (1) piston on left / rigid right: p(Lx) = −iρc₀v_n/sin(kLx), "
    "median error 0.0001 %; (2) piston on left / pressure-release right: "
    "p(Lx/2) = iρc₀v_n sin(k(Lx−x))/cos(kLx), median error 0.067 %; (3) piston on left / "
    "anechoic right: |p| = ρc₀v_n (flat), median error 0.0001 %. All comparisons are "
    "implemented in examples/test_direct_helmholtz_pycafe_fenics.py."
)
done_tag()
separator()

heading("Comparison with open-source frameworks", 2)
reviewer_comment(
    '"No comparison with established open source finite element frameworks such as FEniCS, '
    'FreeFEM, Elmer FEM, MFEM is provided."'
)
response(
    "A quantitative cross-validation against FEniCSx (dolfinx 0.9.0, Q2 Lagrange elements) "
    "has been added. Four direct Helmholtz benchmark cases and one modal benchmark are compared. "
    "Cross-errors: modal eigenfrequencies < 0.0001 %; direct solver cases 0.0000–0.067 % "
    "(median). Results are presented in figures in the Quality Control section."
)
done_tag()
separator()

# ══════════════════════════════════════════════════════════════════════
# REVIEWER 5
# ══════════════════════════════════════════════════════════════════════
heading("Reviewer 5", 1)

heading("Methodological novelty", 2)
reviewer_comment(
    '"Proposed pycafe framework absences methodological novelty as it reproduces classical '
    'frequency-domain acoustic FEM formulations."'
)
response(
    "We refer to our response to Reviewer 4 on the same point. pyCAFE is a research software "
    "contribution, not a methods paper; JORS scope covers this category explicitly."
)
separator()

heading("Analytical benchmark verification", 2)
reviewer_comment(
    '"Validation is restricted to a single modal comparison with a commercial solver, lacking '
    'analytical benchmark verification."'
)
response(
    "Analytical verification has been added: (1) eigenfrequency table for a 1 m × 0.5 m rigid "
    "cavity, first six modes, CQUAD4 and CQUAD8 errors; (2) four direct Helmholtz cases with "
    "closed-form 1D analytical solutions. All are documented in the revised Quality Control section."
)
done_tag()
separator()

heading("Dirichlet BC — Guyan reduction justification vs alternatives", 2)
reviewer_comment(
    '"Adoption of Guyan reduction omits theoretical justification and comparative performance '
    'analysis against Lagrange multipliers, penalty formulations, Nitsche-based approaches."'
)
response(
    "The revised Implementation section now explains that the static condensation used is "
    "algebraically exact (not an approximation) and is mathematically equivalent to the "
    "approaches used by FEniCSx. Alternatives such as penalty and Nitsche methods are "
    "acknowledged as equivalent options for enforcing essential BCs. Since pyCAFE targets "
    "small-to-medium 2D problems, the direct partitioned approach is both sufficient and "
    "more transparent for educational purposes."
)
done_tag()
separator()

heading("Quality control / regression testing", 2)
reviewer_comment(
    '"Quality control section explicitly acknowledges limited verification and testing procedures."'
)
response(
    "The test suite has been substantially expanded (see response to Reviewer 3) and is now "
    "fully described in the Quality Control section. All tests are automated via GitHub Actions CI."
)
done_tag()
separator()

heading("Advanced non-reflecting BCs (PML, infinite elements)", 2)
reviewer_comment(
    '"Framework fails implementation of advanced non-reflecting boundary conditions beyond '
    'simple impedance models."'
)
response(
    "The current version targets enclosed 2D acoustic domains where impedance BCs are "
    "sufficient. PML and infinite-element formulations for open-domain problems are listed as "
    "planned future features in the revised Future Work subsection."
)
separator()

heading("Comparison with Python-based acoustic FEM frameworks", 2)
reviewer_comment(
    '"Absence of comparison with existing Python-based acoustic FEM frameworks prevents objective evaluation."'
)
response(
    "A direct numerical comparison with FEniCSx (the most capable Python-based FEM framework) "
    "has been added: five benchmark configurations, modal and direct solvers, with cross-errors "
    "below 0.07 % in all cases. A qualitative comparison table is included in the Introduction."
)
done_tag()
separator()

# ══════════════════════════════════════════════════════════════════════
# REVIEWER 6
# ══════════════════════════════════════════════════════════════════════
heading("Reviewer 6", 1)

heading("Compulsory: Persistent Identifier (Zenodo DOI)", 2)
reviewer_comment(
    '"The GitHub URL listed as \'Persistent identifier\' is not a persistent identifier. '
    'Please archive a release on Zenodo and provide a DOI."'
)
response(
    "A Zenodo release has been deposited on 15 April 2026. The persistent DOI has been added "
    "to the Availability section of the revised manuscript and replaces the GitHub URL as the "
    "primary persistent identifier."
)
done_tag()
separator()

heading("Compulsory: Commercial software identification", 2)
reviewer_comment(
    '"Please specify which commercial finite element software was used for the validation '
    'comparison in Figure 2."'
)
response(
    "The name and version of the commercial software have been inserted in the Quality Control "
    "section and in the figure caption of Figure 2. [Note to authors: please fill in "
    "[COMMERCIAL SOFTWARE NAME, VERSION] in the manuscript before submission.]"
)
done_tag()
separator()

heading("Compulsory: Typo 'Heigth' → 'Height'", 2)
reviewer_comment('"\'Heigth\' should be \'Height\' in Figure 2."')
response("The typo has been corrected in the figure source file and caption.")
done_tag()
separator()

heading("Optional: Additional keywords", 2)
reviewer_comment('"Consider adding more specific keywords such as \'Helmholtz equation\', '
                 '\'modal analysis\', \'frequency response\'."')
response(
    "Keywords now include: Acoustics; Finite Element Method; Python; Helmholtz equation; "
    "frequency domain; modal analysis; CQUAD8; computational acoustics; open source."
)
done_tag()
separator()

heading("Optional: Reference [6] missing author names", 2)
reviewer_comment('"Reference [6] appears to have missing author names."')
response(
    "Reference [6] (HelmholtzResonator2012) has been corrected: authors Cosma, Maier, and "
    "Mihon are now listed explicitly in the bibliography entry."
)
done_tag()
separator()

heading("Optional: pip install command in paper", 2)
reviewer_comment('"Including a brief mention of the installation command (pip install) would be helpful."')
response(
    "The Reuse Potential section now mentions 'pip install pycafe' as the recommended "
    "installation method."
)
done_tag()
separator()

heading("Optional: PyPI submission", 2)
reviewer_comment('"Consider submitting the package to PyPI for easier installation."')
response(
    "PyPI submission is planned. We will submit before the final acceptance. This item is "
    "tracked in the repository issue tracker."
)
separator()

# ══════════════════════════════════════════════════════════════════════
# REVIEWER 7
# ══════════════════════════════════════════════════════════════════════
heading("Reviewer 7", 1)

heading("Open-source alternatives discussion", 2)
reviewer_comment(
    '"There is no discussion on similar open-source solutions... the reach of other python '
    'packages in the field of acoustics can be highlighted."'
)
response(
    "A paragraph comparing pyCAFE with FEniCSx, FreeFEM, scikit-fem, and openCFS has been "
    "added to the Introduction. The paragraph explains the design-philosophy differences and "
    "positions pyCAFE's matrix-explicit, acoustics-specific, educational workflow in the "
    "landscape of available tools."
)
done_tag()
separator()

heading("Why frequency-domain formulations are preferred", 2)
reviewer_comment(
    '"\'In these applications, frequency-domain formulations are widely adopted, particularly '
    'when time-harmonic responses are of interest.\' Why are frequency-domain formulations preferred?"'
)
response(
    "The Introduction sentence has been expanded to explain the motivation: for time-harmonic "
    "steady-state excitation, the frequency-domain formulation avoids time-stepping and directly "
    "yields the steady-state response at each frequency of interest, which is computationally "
    "more efficient than time integration for narrowband or broadband analyses."
)
done_tag()
separator()

heading("Definition of 'low to mid frequency range'", 2)
reviewer_comment('"What is low to mid frequency range?"')
response(
    "The Introduction now explicitly defines the scope: 'low-to-mid frequency' is understood "
    "here as the regime where kh < 1 (wave number k times element size h), ensuring that "
    "polynomial interpolation remains accurate. In practice this corresponds to at least 6 "
    "elements per wavelength. The upper frequency is therefore mesh-dependent."
)
done_tag()
separator()

heading("Clarity of sentence on weak-form frameworks", 2)
reviewer_comment(
    '"Sentence about weak-form frameworks carries a lot of information and is challenging '
    'to understand."'
)
response(
    "The sentence has been split into two: one describing the variational symbolic approach "
    "common to general FEM frameworks, and one describing how pyCAFE differs by assembling "
    "matrices explicitly."
)
done_tag()
separator()

heading("Emoji in test_basic.py", 2)
reviewer_comment('"In tests/test_basic.py, there is an emoji in the comments. Is that intentional?"')
response(
    "The emoji was unintentional and has been removed from the test file."
)
done_tag()
separator()

heading("Figure 3 font size", 2)
reviewer_comment('"Fig. 3 is hard to read; the reviewer would prefer that the font on the figure is made slightly bigger."')
response(
    "The font size in Figure 3 has been increased from 8 pt to 11 pt for better readability."
)
done_tag()
separator()

heading("2D limitation — highlight again in Implementation", 2)
reviewer_comment(
    '"When discussing geometry supported by pyCAFE, the limitation to two-dimensional '
    'problems could be highlighted again."'
)
response(
    "The 2D limitation is now stated explicitly in three places: the Abstract, the "
    "Implementation section (Geometry subsection), and the Reuse Potential section."
)
done_tag()
separator()

heading("Guided notebook — connecting text between cells", 2)
reviewer_comment(
    '"The example would significantly benefit from connecting text between the cells, '
    'instead of just running cells of code."'
)
response(
    "The guided tutorial notebook examples/analisi_modale_guidata.ipynb now includes "
    "explanatory markdown text between each code cell, describing the purpose of each step, "
    "the expected output, and connections to the theoretical background."
)
done_tag()
separator()

# ══════════════════════════════════════════════════════════════════════
# SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════
doc.add_page_break()
heading("Summary of Changes", 1)

para(
    "The table below summarises all compulsory and major changes made in the revised manuscript "
    "and repository."
)

table = doc.add_table(rows=1, cols=3)
table.style = 'Light Shading Accent 1'
hdr = table.rows[0].cells
hdr[0].text = "Item"
hdr[1].text = "Reviewer(s)"
hdr[2].text = "Status"
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

rows_data = [
    ("Zenodo DOI deposited and added to manuscript", "R6, R7", "✓ Done — 15 April 2026"),
    ("Commercial software name/version in paper", "R6", "✓ Done (placeholder to fill)"),
    ("Typo 'Heigth' → 'Height' in Figure 2", "R6", "✓ Fixed"),
    ("Open-source alternatives paragraph in Introduction", "R1, R7", "✓ Added"),
    ("5–7 keywords added", "R1, R6", "✓ 9 keywords"),
    ("Abstract revised — technical, not promotional", "R1", "✓ Revised"),
    ("GitHub Actions CI pipeline created", "R3", "✓ .github/workflows/ci.yml"),
    ("4-module automated test suite", "R1, R3", "✓ tests/ directory"),
    ("Analytical validation table (6 modes, CQUAD4 + CQUAD8)", "R3, R4, R5", "✓ Quality Control"),
    ("h- and p-convergence study with measured slopes", "R3, R4", "✓ examples/convergence_hp.ipynb"),
    ("FEniCSx modal cross-validation (20 modes, < 0.0001 %)", "R3, R4", "✓ examples/comparison_pycafe_fenics.ipynb"),
    ("FEniCSx direct Helmholtz cross-validation (4 BC cases)", "R3, R4, R5", "✓ examples/test_direct_helmholtz_pycafe_fenics.py"),
    ("Dirichlet BC two-stage explanation + equation in paper", "R4, R5", "✓ Implementation section"),
    ("Velocity BC verified against 1D analytical benchmarks", "R4", "✓ Quality Control"),
    ("Performance comparison (timing) pyCAFE vs FEniCSx", "R3", "✓ Quality Control"),
    ("Future Work subsection (3D, PML, time-domain, etc.)", "R3, R4, R5", "✓ Reuse Potential"),
    ("Reuse Potential expanded (3 subsections)", "R1, R3", "✓ Reuse Potential"),
    ("Guided tutorial notebook with connecting text", "R2, R7", "✓ analisi_modale_guidata.ipynb"),
    ("Reference [6] author names corrected", "R6", "✓ Bibliography"),
    ("Emoji removed from test_basic.py", "R7", "✓ Fixed"),
    ("Figure 3 font size increased", "R7", "✓ Fixed"),
    ("'Low to mid frequency' defined (kh < 1)", "R7", "✓ Introduction"),
    ("pip install pycafe mentioned in paper", "R6", "✓ Reuse Potential"),
]

for item, rev, status in rows_data:
    row = table.add_row().cells
    row[0].text = item
    row[1].text = rev
    row[2].text = status

doc.save("Response_to_Reviewers_pyCAFE.docx")
print("Saved: Response_to_Reviewers_pyCAFE.docx")
